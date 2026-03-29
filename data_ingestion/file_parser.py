"""
Enterprise ETL Platform - Universal File Parser

Parses CSV, Excel (.xlsx), JSON, TXT, PDF, and Word (.docx) files
into a normalised pandas DataFrame with metadata.

Usage:
    result = FileParser.parse(path, file_type="csv")
"""
import io
import json
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd

# ── optional heavy imports (graceful fallback) ────────────────────────────────
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ParseResult:
    """Result of parsing a single uploaded file."""
    filename: str
    file_type: str
    data: pd.DataFrame
    row_count: int = 0
    col_count: int = 0
    file_size_bytes: int = 0
    parse_warnings: List[str] = field(default_factory=list)
    extra_text: Optional[str] = None      # For PDF / DOCX raw text
    sheet_names: Optional[List[str]] = None  # For Excel multi-sheet
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        if self.row_count == 0 and not self.data.empty:
            self.row_count = len(self.data)
            self.col_count = len(self.data.columns)

    @property
    def size_kb(self) -> float:
        return round(self.file_size_bytes / 1024, 1)


SUPPORTED_TYPES = {
    ".csv":  "csv",
    ".tsv":  "csv",
    ".xlsx": "excel",
    ".xls":  "excel",
    ".json": "json",
    ".jsonl":"json",
    ".txt":  "txt",
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx",
}


class FileParser:
    """
    Stateless utility class for parsing all supported file formats
    into a (DataFrame, ParseResult) pair.
    """

    @classmethod
    def parse(
        cls,
        source,                     # path string, Path, or bytes-like / BytesIO
        filename: str = "upload",
        file_type: Optional[str] = None,
        excel_sheet: Optional[str] = None,
    ) -> ParseResult:
        """
        Parse a file and return a ParseResult.

        Args:
            source:      File path or bytes content.
            filename:    Original file name (used to infer type).
            file_type:   Override detected type.
            excel_sheet: Sheet to read from Excel (None = first sheet).

        Returns:
            ParseResult with populated .data DataFrame.
        """
        suffix = Path(filename).suffix.lower()
        detected_type = file_type or SUPPORTED_TYPES.get(suffix, "csv")

        # Normalise source to bytes for consistent handling
        if isinstance(source, (str, Path)):
            raw = Path(source).read_bytes()
            file_size = os.path.getsize(source)
        elif isinstance(source, (bytes, bytearray)):
            raw = bytes(source)
            file_size = len(raw)
        else:
            # BytesIO / UploadedFile
            raw = source.read() if hasattr(source, "read") else bytes(source)
            file_size = len(raw)

        parse_fn = {
            "csv":   cls._parse_csv,
            "excel": cls._parse_excel,
            "json":  cls._parse_json,
            "txt":   cls._parse_txt,
            "pdf":   cls._parse_pdf,
            "docx":  cls._parse_docx,
        }.get(detected_type, cls._parse_csv)

        result = parse_fn(raw, filename, excel_sheet)
        result.file_size_bytes = file_size
        return result

    # ── CSV / TSV ─────────────────────────────────────────────────────────────
    @classmethod
    def _parse_csv(cls, raw: bytes, filename: str, _) -> ParseResult:
        warnings = []
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                sep = "\t" if filename.endswith(".tsv") else ","
                df = pd.read_csv(
                    io.BytesIO(raw),
                    encoding=enc,
                    sep=sep,
                    on_bad_lines="warn",
                    low_memory=False,
                )
                return ParseResult(
                    filename=filename, file_type="csv",
                    data=df, parse_warnings=warnings,
                )
            except UnicodeDecodeError:
                warnings.append(f"Encoding {enc} failed, trying next...")
        raise ValueError(f"Could not decode {filename} with any supported encoding.")

    # ── Excel ─────────────────────────────────────────────────────────────────
    @classmethod
    def _parse_excel(cls, raw: bytes, filename: str, sheet_name) -> ParseResult:
        warnings = []
        xl = pd.ExcelFile(io.BytesIO(raw), engine="openpyxl")
        sheets = xl.sheet_names
        target = sheet_name if sheet_name and sheet_name in sheets else sheets[0]
        if sheet_name and sheet_name not in sheets:
            warnings.append(f"Sheet '{sheet_name}' not found. Loaded '{target}'.")
        df = pd.read_excel(io.BytesIO(raw), sheet_name=target, engine="openpyxl")
        return ParseResult(
            filename=filename, file_type="excel",
            data=df, parse_warnings=warnings,
            sheet_names=sheets,
            metadata={"loaded_sheet": target},
        )

    # ── JSON ──────────────────────────────────────────────────────────────────
    @classmethod
    def _parse_json(cls, raw: bytes, filename: str, _) -> ParseResult:
        warnings = []
        text = raw.decode("utf-8", errors="replace")
        # Try line-delimited first
        if filename.endswith(".jsonl") or "\n{" in text:
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            try:
                records = [json.loads(l) for l in lines]
                df = pd.DataFrame(records)
                return ParseResult(filename=filename, file_type="json", data=df)
            except Exception:
                pass
        try:
            obj = json.loads(text)
            if isinstance(obj, list):
                df = pd.DataFrame(obj)
            elif isinstance(obj, dict):
                # Try common keys that hold records
                for key in ("data", "results", "records", "items", "rows"):
                    if key in obj and isinstance(obj[key], list):
                        df = pd.DataFrame(obj[key])
                        warnings.append(f"Extracted records from JSON key: '{key}'")
                        break
                else:
                    df = pd.json_normalize(obj)
            else:
                df = pd.DataFrame([{"value": obj}])
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON in {filename}: {e}")
        return ParseResult(
            filename=filename, file_type="json",
            data=df, parse_warnings=warnings,
        )

    # ── TXT ───────────────────────────────────────────────────────────────────
    @classmethod
    def _parse_txt(cls, raw: bytes, filename: str, _) -> ParseResult:
        text = raw.decode("utf-8", errors="replace")
        lines = [l.rstrip() for l in text.splitlines() if l.strip()]
        # Try to detect delimited structure
        sample = lines[:5]
        for sep in (",", "\t", "|", ";"):
            if all(sep in l for l in sample):
                try:
                    df = pd.read_csv(io.StringIO(text), sep=sep, on_bad_lines="warn")
                    return ParseResult(
                        filename=filename, file_type="txt", data=df,
                        parse_warnings=[f"Detected delimiter: '{sep}'"],
                    )
                except Exception:
                    pass
        # Fall back: one column of text lines
        df = pd.DataFrame({"line_number": range(1, len(lines) + 1), "text": lines})
        return ParseResult(
            filename=filename, file_type="txt",
            data=df, extra_text=text,
            parse_warnings=["No delimiter detected — loaded as line-by-line text."],
        )

    # ── PDF ───────────────────────────────────────────────────────────────────
    @classmethod
    def _parse_pdf(cls, raw: bytes, filename: str, _) -> ParseResult:
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")

        all_tables: List[pd.DataFrame] = []
        all_text_lines: List[str] = []

        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                # Extract tables
                for table in page.extract_tables():
                    if table and len(table) > 1:
                        header = [str(c) if c else f"col_{i}" for i, c in enumerate(table[0])]
                        rows = table[1:]
                        all_tables.append(pd.DataFrame(rows, columns=header))
                # Extract text
                text = page.extract_text() or ""
                all_text_lines.extend(text.splitlines())

        full_text = "\n".join(all_text_lines)

        if all_tables:
            df = pd.concat(all_tables, ignore_index=True)
        else:
            # No tables found — represent text as structured rows
            df = pd.DataFrame({
                "line": range(1, len(all_text_lines) + 1),
                "content": all_text_lines,
            })

        return ParseResult(
            filename=filename, file_type="pdf",
            data=df, extra_text=full_text,
            metadata={"pages": total_pages, "tables_found": len(all_tables)},
            parse_warnings=[] if all_tables else ["No tables found; loaded as text lines."],
        )

    # ── DOCX ──────────────────────────────────────────────────────────────────
    @classmethod
    def _parse_docx(cls, raw: bytes, filename: str, _) -> ParseResult:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = DocxDocument(io.BytesIO(raw))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        # Extract tables if present
        table_dfs = []
        for table in doc.tables:
            if not table.rows:
                continue
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])
            if headers and rows:
                table_dfs.append(pd.DataFrame(rows, columns=headers))

        if table_dfs:
            df = pd.concat(table_dfs, ignore_index=True)
        else:
            df = pd.DataFrame({
                "paragraph": range(1, len(paragraphs) + 1),
                "content": paragraphs,
            })

        return ParseResult(
            filename=filename, file_type="docx",
            data=df, extra_text=full_text,
            metadata={"paragraphs": len(paragraphs), "tables": len(table_dfs)},
            parse_warnings=[] if table_dfs else ["No tables found; loaded paragraphs as rows."],
        )
